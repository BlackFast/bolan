#!/usr/bin/env python3
"""Extract plain text from a book file and split it into markdown chunks."""

from __future__ import annotations

import argparse
import html
import json
import re
import sys
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable


TEXT_SUFFIXES = {".txt", ".md", ".markdown", ".rst"}


@dataclass
class SourceResult:
    text: str
    warnings: list[str]


def read_text_file(path: Path) -> SourceResult:
    warnings: list[str] = []
    for encoding in ("utf-8", "utf-8-sig", "gb18030", "latin-1"):
        try:
            return SourceResult(path.read_text(encoding=encoding), warnings)
        except UnicodeDecodeError:
            continue
    raise RuntimeError(f"Could not decode text file: {path}")


def strip_html(value: str) -> str:
    value = re.sub(r"(?is)<(script|style).*?>.*?</\1>", " ", value)
    value = re.sub(r"(?s)<[^>]+>", " ", value)
    return html.unescape(value)


def read_docx(path: Path) -> SourceResult:
    try:
        import docx  # type: ignore
    except ImportError as exc:
        raise RuntimeError("DOCX extraction requires python-docx. Install with: pip install python-docx") from exc

    document = docx.Document(str(path))
    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return SourceResult("\n\n".join(parts), [])


def read_pdf(path: Path) -> SourceResult:
    try:
        from pypdf import PdfReader  # type: ignore
    except ImportError:
        try:
            from PyPDF2 import PdfReader  # type: ignore
        except ImportError as exc:
            raise RuntimeError("PDF extraction requires pypdf or PyPDF2. Install with: pip install pypdf") from exc

    reader = PdfReader(str(path))
    warnings: list[str] = []
    pages: list[str] = []
    for index, page in enumerate(reader.pages, start=1):
        try:
            text = page.extract_text() or ""
        except Exception as exc:  # pragma: no cover - depends on PDF internals
            warnings.append(f"Page {index}: extraction failed: {exc}")
            text = ""
        if text.strip():
            pages.append(f"\n\n[Page {index}]\n{text}")
    return SourceResult("\n".join(pages), warnings)


def read_epub(path: Path) -> SourceResult:
    warnings: list[str] = []
    try:
        from ebooklib import epub  # type: ignore
    except ImportError:
        return read_epub_zip(path)

    book = epub.read_epub(str(path))
    parts: list[str] = []
    for item in book.get_items():
        media_type = getattr(item, "media_type", "")
        if media_type in {"application/xhtml+xml", "text/html"}:
            parts.append(strip_html(item.get_content().decode("utf-8", errors="ignore")))
    return SourceResult("\n\n".join(parts), warnings)


def read_epub_zip(path: Path) -> SourceResult:
    warnings = ["ebooklib is not installed; used basic EPUB zip extraction instead."]
    parts: list[str] = []
    with zipfile.ZipFile(path) as archive:
        names = sorted(
            name
            for name in archive.namelist()
            if name.lower().endswith((".html", ".xhtml", ".htm"))
        )
        for name in names:
            raw = archive.read(name).decode("utf-8", errors="ignore")
            text = strip_html(raw)
            if text.strip():
                parts.append(f"\n\n[EPUB item: {name}]\n{text}")
    return SourceResult("\n\n".join(parts), warnings)


def read_folder(path: Path) -> SourceResult:
    warnings: list[str] = []
    parts: list[str] = []
    files = sorted(p for p in path.rglob("*") if p.is_file() and p.suffix.lower() in TEXT_SUFFIXES)
    if not files:
        raise RuntimeError(f"No text or markdown files found under folder: {path}")
    for file_path in files:
        result = read_text_file(file_path)
        warnings.extend(f"{file_path}: {warning}" for warning in result.warnings)
        parts.append(f"\n\n[File: {file_path.relative_to(path)}]\n{result.text}")
    return SourceResult("\n".join(parts), warnings)


def read_source(path: Path) -> SourceResult:
    if path.is_dir():
        return read_folder(path)

    suffix = path.suffix.lower()
    if suffix in TEXT_SUFFIXES:
        return read_text_file(path)
    if suffix == ".pdf":
        return read_pdf(path)
    if suffix == ".docx":
        return read_docx(path)
    if suffix == ".epub":
        return read_epub(path)
    raise RuntimeError(f"Unsupported file type: {suffix or '<none>'}")


def normalize_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{4,}", "\n\n\n", text)
    return text.strip()


def split_paragraphs(text: str) -> list[str]:
    paragraphs = [part.strip() for part in re.split(r"\n\s*\n", text) if part.strip()]
    if len(paragraphs) <= 1:
        paragraphs = [part.strip() for part in re.split(r"(?<=[.!?。！？])\s+", text) if part.strip()]
    return paragraphs


def make_chunks(text: str, chunk_chars: int, overlap_chars: int) -> list[str]:
    paragraphs = split_paragraphs(text)
    chunks: list[str] = []
    current: list[str] = []
    current_len = 0

    for paragraph in paragraphs:
        paragraph_len = len(paragraph)
        if current and current_len + paragraph_len + 2 > chunk_chars:
            chunks.append("\n\n".join(current).strip())
            if overlap_chars > 0:
                overlap: list[str] = []
                overlap_len = 0
                for previous in reversed(current):
                    if overlap_len + len(previous) > overlap_chars:
                        break
                    overlap.insert(0, previous)
                    overlap_len += len(previous) + 2
                current = overlap
                current_len = overlap_len
            else:
                current = []
                current_len = 0
        current.append(paragraph)
        current_len += paragraph_len + 2

    if current:
        chunks.append("\n\n".join(current).strip())
    return chunks


def write_chunks(chunks: Iterable[str], chunks_dir: Path) -> list[dict[str, object]]:
    chunks_dir.mkdir(parents=True, exist_ok=True)
    records: list[dict[str, object]] = []
    for index, chunk in enumerate(chunks, start=1):
        filename = f"chunk-{index:04d}.md"
        path = chunks_dir / filename
        path.write_text(f"# Chunk {index:04d}\n\n{chunk}\n", encoding="utf-8")
        records.append({"index": index, "file": str(path), "chars": len(chunk)})
    return records


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("source", type=Path, help="Book file or folder to extract.")
    parser.add_argument("--out", type=Path, required=True, help="Output workspace directory.")
    parser.add_argument("--chunk-chars", type=int, default=12000, help="Target chunk size in characters.")
    parser.add_argument("--overlap-chars", type=int, default=800, help="Paragraph overlap between chunks.")
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    source = args.source.expanduser().resolve()
    out_dir = args.out.expanduser().resolve()

    if args.chunk_chars < 2000:
        raise RuntimeError("--chunk-chars should be at least 2000")
    if args.overlap_chars < 0 or args.overlap_chars >= args.chunk_chars:
        raise RuntimeError("--overlap-chars must be non-negative and smaller than --chunk-chars")
    if not source.exists():
        raise RuntimeError(f"Source does not exist: {source}")

    result = read_source(source)
    text = normalize_text(result.text)
    if not text:
        raise RuntimeError("No extractable text found.")

    text_dir = out_dir / "text"
    chunks_dir = out_dir / "chunks"
    text_dir.mkdir(parents=True, exist_ok=True)
    text_path = text_dir / "book.txt"
    text_path.write_text(text + "\n", encoding="utf-8")

    chunks = make_chunks(text, args.chunk_chars, args.overlap_chars)
    chunk_records = write_chunks(chunks, chunks_dir)

    manifest = {
        "source": str(source),
        "output": str(out_dir),
        "text_file": str(text_path),
        "characters": len(text),
        "chunks": chunk_records,
        "warnings": result.warnings,
    }
    manifest_path = out_dir / "manifest.json"
    manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")

    print(f"Wrote text: {text_path}")
    print(f"Wrote chunks: {chunks_dir} ({len(chunk_records)} files)")
    print(f"Wrote manifest: {manifest_path}")
    for warning in result.warnings:
        print(f"Warning: {warning}", file=sys.stderr)
    return 0


if __name__ == "__main__":
    try:
        raise SystemExit(main())
    except Exception as exc:
        print(f"Error: {exc}", file=sys.stderr)
        raise SystemExit(1)
